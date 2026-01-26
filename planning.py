from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, timedelta
import calendar


# =====================================================
# MOIS ET JOURS EN FRANÇAIS
# =====================================================

MOIS_FR = [
    "", "JANVIER", "FÉVRIER", "MARS", "AVRIL", "MAI", "JUIN",
    "JUILLET", "AOÛT", "SEPTEMBRE", "OCTOBRE", "NOVEMBRE", "DÉCEMBRE"
]

JOURS_FR = ["Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche"]


# =====================================================
# COULEURS MODERNES
# =====================================================

COLORS = {
    "ENTREPRISE": "FFF4CC",
    "FORMATION": "DCEAF7",
    "FERIE": "FFD6D6",
    "WEEKEND": "F2F2F2",
}


# =====================================================
# FORMATION (modifiable)
# =====================================================

RANGES_FORMATION = [
    (2026, 1, 19, 23),
    (2026, 2, 9, 13),
    (2026, 3, 2, 6),
    (2026, 3, 23, 27),
    (2026, 4, 13, 17),
    (2026, 5, 4, 7),
    (2026, 5, 26, 29),
    (2026, 6, 15, 19),
    (2026, 6, 22, 26),
    (2026, 6, 29, 30),
    (2026, 7, 1, 3),
    (2026, 7, 6, 10),
    (2026, 9, 14, 18),
]


# =====================================================
# JOURS FÉRIÉS FRANCE
# =====================================================

def paques(annee):
    a = annee % 19
    b = annee // 100
    c = annee % 100
    d = b // 4
    e = b % 4
    f = (b + 8) // 25
    g = (b - f + 1) // 3
    h = (19*a + b - d - g + 15) % 30
    i = c // 4
    k = c % 4
    l = (32 + 2*e + 2*i - h - k) % 7
    m = (a + 11*h + 22*l) // 451
    mois = (h + l - 7*m + 114) // 31
    jour = (h + l - 7*m + 114) % 31 + 1
    return date(annee, mois, jour)


def jours_feries_france(annee):
    p = paques(annee)
    return {
        date(annee, 1, 1),
        p + timedelta(days=1),
        date(annee, 5, 1),
        date(annee, 5, 8),
        p + timedelta(days=39),
        p + timedelta(days=50),
        date(annee, 7, 14),
        date(annee, 8, 15),
        date(annee, 11, 1),
        date(annee, 11, 11),
        date(annee, 12, 25),
    }


# =====================================================
# TYPE DE JOUR
# =====================================================

def get_day_type(d):
    if d in jours_feries_france(d.year):
        return "FERIE"
    if d.weekday() >= 5:
        return "WEEKEND"

    for y, m, d1, d2 in RANGES_FORMATION:
        if d.year == y and d.month == m and d1 <= d.day <= d2:
            return "FORMATION"

    return "ENTREPRISE"


# =====================================================
# COULEUR CELLULE WORD
# =====================================================

def set_cell_color(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


# =====================================================
# CRÉATION D’UN MOIS
# =====================================================

def creer_mois(doc, annee, mois):

    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    titre = doc.add_heading(f"{MOIS_FR[mois]} {annee}", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"

    for i, j in enumerate(JOURS_FR):
        table.rows[0].cells[i].text = j

    cal = calendar.monthcalendar(annee, mois)

    for semaine in cal:
        row = table.add_row().cells
        for col, jour in enumerate(semaine):
            if jour == 0:
                continue

            d = date(annee, mois, jour)
            dtype = get_day_type(d)

            cell = row[col]
            cell.text = str(jour)
            set_cell_color(cell, COLORS[dtype])

    doc.add_page_break()


# =====================================================
# GÉNÉRATION 2026
# =====================================================

def generer_word(filename):
    doc = Document()
    doc.add_heading("CALENDRIER ALTERNANCE 2026", 0)

    for mois in range(1, 13):
        creer_mois(doc, 2026, mois)

    doc.save(filename)
    print("Fichier généré :", filename)


if __name__ == "__main__":
    generer_word("calendrier_2026_francais.docx")
